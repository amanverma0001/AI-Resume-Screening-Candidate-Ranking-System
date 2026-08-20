"""
File Parsers for PDF, DOCX, and TXT resumes and job descriptions.
Supports both file paths and in-memory UploadedFile streams.
"""

import io
from typing import Union, BinaryIO
from core.text_cleaner import clean_text

def parse_pdf(file_input: Union[str, BinaryIO, bytes]) -> str:
    """Extracts text from a PDF file path or bytes buffer."""
    text_content = []
    try:
        from pypdf import PdfReader
        
        if isinstance(file_input, (bytes, bytearray)):
            stream = io.BytesIO(file_input)
        elif hasattr(file_input, 'read'):
            stream = file_input
        else:
            stream = open(file_input, 'rb')
            
        reader = PdfReader(stream)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
                
    except Exception as e:
        print(f"[PDF Parser Error] {e}")
        
    return clean_text("\n".join(text_content))

def parse_docx(file_input: Union[str, BinaryIO, bytes]) -> str:
    """Extracts text from a DOCX Word document."""
    text_content = []
    try:
        import docx
        
        if isinstance(file_input, (bytes, bytearray)):
            stream = io.BytesIO(file_input)
        elif hasattr(file_input, 'read'):
            stream = file_input
        else:
            stream = open(file_input, 'rb')
            
        doc = docx.Document(stream)
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
    except Exception as e:
        print(f"[DOCX Parser Error] {e}")
        
    return clean_text("\n".join(text_content))

def parse_txt(file_input: Union[str, BinaryIO, bytes]) -> str:
    """Extracts text from TXT or raw byte streams."""
    try:
        if isinstance(file_input, (bytes, bytearray)):
            return clean_text(file_input.decode('utf-8', errors='ignore'))
        elif hasattr(file_input, 'read'):
            data = file_input.read()
            if isinstance(data, bytes):
                return clean_text(data.decode('utf-8', errors='ignore'))
            return clean_text(str(data))
        else:
            with open(file_input, 'r', encoding='utf-8', errors='ignore') as f:
                return clean_text(f.read())
    except Exception as e:
        print(f"[TXT Parser Error] {e}")
        return ""

def extract_document_text(file_input: Union[str, BinaryIO, bytes], filename: str = "") -> str:
    """Universal dispatcher that detects file format and parses accordingly."""
    name_lower = filename.lower()
    
    if name_lower.endswith('.pdf'):
        return parse_pdf(file_input)
    elif name_lower.endswith('.docx'):
        return parse_docx(file_input)
    elif name_lower.endswith('.txt') or name_lower.endswith('.md'):
        return parse_txt(file_input)
    else:
        # Try PDF first, then fallback to TXT
        text = parse_pdf(file_input)
        if not text:
            text = parse_docx(file_input)
        if not text:
            text = parse_txt(file_input)
        return text
